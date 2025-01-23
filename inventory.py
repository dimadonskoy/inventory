from flask import Flask, render_template, request
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from docx.shared import Pt
import os
import sys
from docx.shared import Inches
import base64

### Working directory ###
if getattr(sys, 'frozen', False):
    workdir = os.path.dirname(sys.executable)
else:
    workdir = os.path.dirname(os.path.abspath(__file__))


inventory = Flask(__name__)


@inventory.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        name = request.form['name']
        input1 = request.form['input1']
        input2 = request.form['input2']
        input3 = request.form['input3']
        input4 = request.form['input4']
        input5 = request.form['input5']
        input6 = request.form['input6']
        input7 = request.form['input7']
        input8 = request.form['input8']
        input9 = request.form['input9']

        # Create a new document
        document = Document()
        # set RTL direction for the document
        document.core_properties.language = 'he-IL'  # Set to Hebrew (Israel)
        # Replace with a Hebrew font name
        document.styles['Normal'].font.name = 'Tahoma'
        document.styles['Normal'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        document.add_heading("ספירת מלאי רכב\n")
        document.element.rtl = True

        # Add a header with the current date and the name
        header = document.sections[0].header

        # Add logo to the header
        logo_path = os.path.join(workdir, 'pics', 'logo.png')
        logo_width = Inches(1.9)
        logo_height = Inches(0.4)
        logo_paragraph = header.add_paragraph()
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=logo_width, height=logo_height)
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the current date to the header
        date_paragraph = header.add_paragraph()
        date_paragraph.text = datetime.now().strftime(
            '%d.%m.%Y')  # Set the text to the current date
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add the name to the header
        date_paragraph = header.add_paragraph()
        date_paragraph.text = name
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add the inputs to the document
        table = document.add_table(rows=9, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'  # add table style

        for row in table.rows:
            for cell in row.cells:
                cell = table.cell(0, 0)
                cell_text = "Text in cell"
                paragraph = cell.add_paragraph(cell_text)
                paragraph.style = document.styles['Normal']
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                paragraph.runs[0].font.size = Pt(16)

        table.cell(0, 0).text = input1
        table.cell(0, 1).text = 'סורק ברקוד חדש'
        table.cell(1, 0).text = input2
        table.cell(1, 1).text = 'FingerPrint'
        table.cell(2, 0).text = input3
        table.cell(2, 1).text = 'VX820'
        table.cell(3, 0).text = input4
        table.cell(3, 1).text = 'Network switch'
        table.cell(4, 0).text = input5
        table.cell(4, 1).text = 'USB מקלדת'
        table.cell(5, 0).text = input6
        table.cell(5, 1).text = 'ללא סים RUT950 נתב'
        table.cell(6, 0).text = input7
        table.cell(6, 1).text = 'עם סים פעיל RUT950 נתב'
        table.cell(7, 0).text = input8
        table.cell(7, 1).text = 'P400'
        table.cell(8, 0).text = input9
        table.cell(8, 1).text = 'מדפסת'

        ##############################  SIGNATURE ########################
        # Create a title paragraph for the signature
        title = document.add_paragraph(style='Normal')
        title.add_run('\n\n\n\n:חתימת המתקין')
        file_path = os.path.join(workdir, 'files', 'signature.png')
        # Save signature as image file
        with open(file_path, 'wb') as f:
            signature_data = request.form['signature_data']
            signature_data = signature_data.replace(
                'data:image/png;base64,', '')
            f.write(base64.b64decode(signature_data))

        # Add signature image to document
        document.add_picture(file_path, width=Inches(2.0))

        # Save the document
        filename = name + '_' + datetime.now().strftime('%d-%m-%Y') + '.docx'
        document.save(workdir + '//files//' + filename)
        filename_without_path = name + '_' + datetime.now().strftime('%d-%m-%Y') + '.docx'

        #########################################   Send email #############################################

        import email
        import smtplib
        import ssl
        from email import encoders
        from email.mime.base import MIMEBase
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText
        from email.message import EmailMessage

        subject = f"טופס ספירת מלאי רכב ({name})"
        body = "VICO Field services team"
        sender_email = "vrics@verifone.co.il"
        receiver_email = ["dimitry_d1@VERIFONE.com"]
        # receiver_email = ["dimitry_d1@VERIFONE.com","i.vird.pos@verifone.com"]

        # Create a multipart message and set headers
        message = MIMEMultipart()
        message["From"] = "VIRD installation team<{}>".format(sender_email)
        message["To"] = ', '.join(receiver_email)
        message["Subject"] = subject
        message.attach(MIMEText(body, "plain"))

        filename = os.path.join(workdir, "files", filename)
        filename_without_path = name + '_' + datetime.now().strftime('%d-%m-%Y') + '.docx'

        # Open the file in binary mode and read its contents
        with open(filename, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)

        # Encode the filename with UTF-8
        filename_encoded = filename_without_path.encode("utf-8")

        # Set the Content-Disposition header with the encoded filename
        part.add_header(
            "Content-Disposition",
            "attachment",
            filename=("utf-8", "", filename_encoded.decode("utf-8"))
        )

        message.attach(part)
        text = message.as_string()

        with smtplib.SMTP("verifonecp.agas.co.il", port=587) as server:
            server.starttls()
            server.login(sender_email, "(K$6&Mu($^cB")
            server.sendmail(sender_email, receiver_email, text)

        ### remove temporary file ###
        if os.path.exists(filename):
            os.remove(filename)

        ################################################################################

        return '''<html>
        <head>
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
        </head>
        <body>
            <div style="
                                display: flex;
                                justify-content: center;
                                align-items: center;
                                height: 100%;
                                font-size: 32px;
                                font-weight: bold;
                            ">
            File sent by email to BO VIRD {}
            </div>
        </body>
        </html>
        '''.format(filename_without_path)

    return render_template('index.html')


if __name__ == '__main__':
    inventory.run(host='0.0.0.0', port=5000)
